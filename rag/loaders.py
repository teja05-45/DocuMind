from __future__ import annotations
from typing import List
import os
import re
import requests
from bs4 import BeautifulSoup

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument


def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text


def load_pdf(file_path: str) -> List[Document]:
    loader = PyPDFLoader(file_path)
    docs = loader.load()
    for d in docs:
        d.page_content = clean_text(d.page_content)
        d.metadata["source"] = os.path.basename(file_path)
        d.metadata["type"] = "pdf"
    return docs


def load_docx(file_path: str) -> List[Document]:
    doc = DocxDocument(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    full_text = clean_text(full_text)

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": os.path.basename(file_path),
                "type": "docx",
            },
        )
    ]


def load_txt(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    text = clean_text(text)

    return [
        Document(
            page_content=text,
            metadata={
                "source": os.path.basename(file_path),
                "type": "txt",
            },
        )
    ]


def load_url(url: str, timeout: int = 10) -> List[Document]:
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) RAG-Bot/1.0"}
    resp = requests.get(url, headers=headers, timeout=timeout)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "lxml")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    text = soup.get_text(separator=" ")
    text = clean_text(text)

    return [
        Document(
            page_content=text,
            metadata={
                "source": url,
                "type": "url",
            },
        )
    ]


def load_notes(notes: str, title: str = "User Notes") -> List[Document]:
    notes = clean_text(notes)
    if not notes:
        return []

    return [
        Document(
            page_content=notes,
            metadata={
                "source": title,
                "type": "notes",
            },
        )
    ]
